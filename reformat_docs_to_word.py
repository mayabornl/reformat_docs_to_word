import os
from pydoc import text
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from lxml import etree

# set relative directory

dirname = os.path.dirname(__file__)
filename_1 = os.path.join(dirname, 'changed_doc/')
filename_2 = os.path.join(dirname,'original_doc/')

# pip install python-docx for the package

def replacing_format(x):
    print(x)
    _, tail = os.path.split(x)
    document = Document(x)

    strikethroughs = []

    for paragraph in document.paragraphs:
        strikethrough = ""

        for run in paragraph.runs:
            if run.font.strike:
                # THESE next few if statements are to fix that odd random text that is striken through
                if '\n' in run.text:
                    run.font.strike = False
                    continue
                if ', please provide a URL to the dataset(s)' in run.text:
                    run.font.strike = False
                    continue
                if 'yes' in run.text:
                    run.font.strike = False
                    continue

                #This fixes the issue were hyperlinks were ignored
                if paragraph._p.xpath("child::w:hyperlink"):
                    p = paragraph._p
                    keep = p.xpath("descendant::text()")
                    s = ''.join(str(t) for t in keep)
                    
                    paragraph.clear()
                    run = paragraph.add_run(s)
                    
                    #print(s)
                    
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.strike = False
                run.font.underline = WD_UNDERLINE.THICK
                strikethrough += run.text

        if strikethrough:
            strikethroughs.append(strikethrough)
            for s in strikethroughs:
                print(s)

    document.save(filename_1 + tail)


def reformat_run():
    listdir = os.listdir(filename_2)

    for file in listdir:
        if file.endswith(".docx"):
            replacing_format(filename_2 + file)


if __name__ == "__main__":
    reformat_run()